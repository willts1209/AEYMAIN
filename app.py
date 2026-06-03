"""AEY — Voice-to-Table MVP web app.

Run locally:
    py app.py

Run with gunicorn (production / Docker):
    gunicorn -w 1 -b 0.0.0.0:5000 app:app --timeout 600

Required env vars:
    ANTHROPIC_API_KEY  Claude API key (https://console.anthropic.com)
    GROQ_API_KEY       Groq API key  (https://console.groq.com)

External dep: ffmpeg on PATH (used to compress audio before sending to Groq).
"""
from __future__ import annotations
import io
import json
import os
import subprocess
import tempfile

from anthropic import Anthropic
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt
from flask import Flask, jsonify, render_template, request, send_file, send_from_directory
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from groq import Groq
from werkzeug.middleware.proxy_fix import ProxyFix

# ── Prompts ──────────────────────────────────────────────────────────
UK_BAT_PROMPT = (
    "Field notes spoken by a UK ecologist conducting a bat dusk emergence survey. "
    "The speaker describes bat movements, sightings, times in HH:MM format, "
    "building sides, compass directions, and behaviour. "
    "Numbers and times are spoken naturally (e.g. 'twenty-one thirty-three')."
)
# Species names are deliberately NOT in this prompt. Listing them primes
# Whisper to emit those words in low-signal audio (rain, silence, mumbling),
# which produces phantom species in noisy real-world surveys. Whisper is
# good enough at UK bat names natively; occasional mis-spellings of an
# actually-said species are far less harmful than hallucinations.

PARSE_INSTRUCTIONS = """You are parsing a UK bat dusk-survey voice transcript into structured rows.

THE THREE LAWS — in priority order. Higher-numbered laws override lower ones.

LAW 1 — Default to empty.
Every field defaults to "" unless the surveyor states it clearly. Empty is ALWAYS a valid answer.
- "A bat" / "bat" / "another one" → species is "".
- Speculative species mentions ("could be a Soprano", "Daubies maybe", "sounded like a Nathusius", "looked like maybe a brown long-eared", "called out as Soprano", "mentioned Nathusius") → species is "" and the guess goes in notes (e.g. "Possibly Soprano pipistrelle").
- Only populate species when the surveyor uses a definite identifying verb: "Common pip emerged", "saw a Soprano foraging", "Noctule flew past", "identified as Daubenton's", or similar unambiguous confirmation.

LAW 2 — Honour the surveyor's own disclaimers, globally.
If the transcript contains a global ID disclaimer such as:
  - "didn't get any confirmed IDs"
  - "no positive IDs"
  - "couldn't tell"
  - "couldn't ID them"
  - "no calling so couldn't ID"
  - "weren't able to confirm"
  - "not sure of any species"
then ALL rows in the output must have species "" regardless of any earlier speculative mentions. Capture movements, locations, directions, and notes as normal, but the species column is empty across the entire survey. Add a single row at the end with the disclaimer in notes if useful, with species/count/behaviour/location/direction left empty.

LAW 3 — Don't invent.
Never add information not in the transcript. Do not annotate your own reasoning ("Species called out during survey", "Mentioned near end") — if the species mention was speculative, just put the guess in notes and leave species empty.

Per-field rules:
- One row per distinct observation explicitly described. Split combined statements.
- time: HH:MM 24-hour if mentioned, else "". For dusk emergence and activity surveys (the default for this app), all observation times fall in the evening (UK ~19:00-23:59). Whisper transcribes spoken evening times in shortened form or with dropped digits ("8:45" for "20:45", "9:43" for "21:43", "1:52" when the leading "2" was dropped from "21:52"). Normalise every parsed time into evening format:
    * HH 8-12 → add 12 (e.g. "8:45" → "20:45", "10:20" → "22:20", "12:05" → "00:05")
    * HH 1-7 → likely a dropped leading "2"; interpret as 21:MM-22:MM (e.g. "1:52" → "21:52", "2:30" → "22:30", "3:15" → "23:15")
    * HH 19-23 → keep as is
    * Use chronological context if ambiguous: observations later in the transcript should generally have later or equal times.
- species: see LAW 1 and LAW 2. Abbreviations expand only when used confidently: CP/Pip = Common pipistrelle, SP = Soprano pipistrelle, NP = Nathusius pipistrelle, NO = Noctule, BLE = Brown long-eared. When in doubt: "".
- count: number as string if mentioned (e.g. "5-8", "2"), else "1".
- behaviour: one of Foraging, Commuting, Flying, Emerged, Re-entry, Social, Other. "Other" if the transcript only describes movement.
- location: building side / vantage point / feature where the observation happened. "" if not mentioned.
- direction: compass direction if mentioned, else "".
- notes: anything else — uncertainty markers ("not calling", "no echolocation heard"), hedged species guesses ("Possibly Soprano pipistrelle"), count caveats, environmental conditions, roost evidence.
- Meta lines (survey start/end times, time-only utterances) → a row with the detail in notes and species/count/behaviour/location/direction all "".
- Roost evidence (emerged from gap/arch/tile/soffit/etc.) → prefix notes with "Possible roost feature: ".

If the transcript contains no actual bat observations at all (pure silence, weather chat, all calls heard but no behavioural detail), return [].

Return ONLY a JSON array. No prose, no markdown fences."""

# ── App ──────────────────────────────────────────────────────────────
app = Flask(__name__)

# Render terminates TLS at its load balancer and forwards via X-Forwarded-For;
# without ProxyFix every request would appear to come from the proxy and the
# rate limiter would key everyone to the same bucket.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1)

CLAUDE = Anthropic()  # ANTHROPIC_API_KEY
GROQ = Groq()         # GROQ_API_KEY

# Rate limiter — protects expensive Groq + Claude calls from accidental hammering
# and adversarial abuse. Storage is in-memory; fine while running on a single
# gunicorn worker (per Dockerfile). Move to redis:// if we ever scale workers.
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=["120 per hour", "600 per day"],
    storage_uri="memory://",
    strategy="fixed-window",
)


def _global_key() -> str:
    """Shared bucket key for global caps on the most expensive endpoints."""
    return "_global_"


@app.errorhandler(429)
def ratelimit_handler(e):
    # The description includes which limit was hit ("10 per 1 hour" etc).
    detail = str(e.description) if hasattr(e, "description") else ""
    return jsonify({
        "error": (
            "Rate limit reached. Wait a few minutes and try again. "
            "If you're a real user hitting this regularly, get in touch: "
            "willts1209@gmail.com"
        ),
        "limit": detail,
    }), 429


def get_audio_duration(path: str) -> float:
    """Audio duration in seconds via ffprobe (ships with ffmpeg)."""
    result = subprocess.run(
        [
            "ffprobe", "-i", path,
            "-show_entries", "format=duration",
            "-v", "quiet", "-of", "csv=p=0",
        ],
        check=True, capture_output=True, text=True,
    )
    return float(result.stdout.strip() or "0")


def compress_audio(input_path: str, output_path: str) -> None:
    """Strip silences, downmix to mono, compress to fit Groq's 25 MB limit.

    The silenceremove filter drops gaps of ≥0.5s below -40 dB. In a typical
    UK bat dusk survey this removes 50-70% of the audio (long quiet stretches
    between observations), which slashes upload time, transcription cost,
    and Whisper drift risk for long recordings.

    Speech quality at mono 32 kbps / 16 kHz is fine — Whisper's native rate.
    """
    subprocess.run(
        [
            "ffmpeg", "-y", "-i", input_path,
            "-af", "silenceremove=stop_periods=-1:stop_duration=0.5:stop_threshold=-40dB",
            "-ac", "1",
            "-b:a", "32k",
            "-ar", "16000",
            output_path,
        ],
        check=True,
        capture_output=True,
    )


GROQ_MAX_BYTES = 25 * 1024 * 1024  # Groq's hard limit on /audio/transcriptions


# ── Routes ───────────────────────────────────────────────────────────
@app.route("/")
@limiter.exempt
def landing():
    return render_template("landing.html")


@app.route("/app")
@limiter.exempt
def index():
    return render_template("index.html")


@app.route("/static/<path:filename>")
@limiter.exempt
def static_files(filename):
    return send_from_directory("static", filename)


@app.route("/healthz")
@limiter.exempt
def healthz():
    return "ok"


@app.route("/api/signup", methods=["POST"])
@limiter.limit("5 per minute")
@limiter.limit("30 per day")
def signup():
    """Early-access email signup. Logs to stdout (visible in Render logs).
    Upgrade to a real datastore once volume warrants it.
    """
    data = request.json or {}
    email = (data.get("email") or "").strip()
    if not email or "@" not in email:
        return jsonify({"error": "invalid email"}), 400
    # stdout shows up in Render's live logs; harvest from there for now.
    app.logger.info("[SIGNUP] %s", email)
    print(f"[SIGNUP] {email}", flush=True)
    return jsonify({"ok": True})


@app.route("/api/transcribe", methods=["POST"])
@limiter.limit("10 per hour")
@limiter.limit("30 per day")
@limiter.limit("100 per day", key_func=_global_key)  # global cap on Groq spend
def transcribe():
    audio = request.files["audio"]
    suffix = os.path.splitext(audio.filename)[1] or ".m4a"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        audio.save(f.name)
        raw_path = f.name
    compressed_path = raw_path + ".compressed.m4a"

    try:
        original_seconds = get_audio_duration(raw_path)
        compress_audio(raw_path, compressed_path)
        compressed_seconds = get_audio_duration(compressed_path)
        compressed_size = os.path.getsize(compressed_path)

        if compressed_size > GROQ_MAX_BYTES:
            return jsonify({
                "error": (
                    f"Audio is still {compressed_size / 1024 / 1024:.1f} MB after silence stripping. "
                    "Groq's hard limit is 25 MB. Chunked transcription is the next feature; for now, "
                    "split the recording in two and process each half separately."
                ),
            }), 413

        with open(compressed_path, "rb") as f:
            transcription = GROQ.audio.transcriptions.create(
                file=(os.path.basename(compressed_path), f.read()),
                model="whisper-large-v3-turbo",
                prompt=UK_BAT_PROMPT,
                response_format="text",
                language="en",
            )

        text = transcription if isinstance(transcription, str) else transcription.text
        return jsonify({
            "transcript": text.strip(),
            "original_seconds": original_seconds,
            "compressed_seconds": compressed_seconds,
            "compressed_bytes": compressed_size,
        })
    finally:
        for p in (raw_path, compressed_path):
            try:
                os.unlink(p)
            except OSError:
                pass


@app.route("/api/parse", methods=["POST"])
@limiter.limit("15 per hour")
@limiter.limit("50 per day")
@limiter.limit("150 per day", key_func=_global_key)  # global cap on Claude spend
def parse():
    transcript = request.json["transcript"]
    user_msg = f"{PARSE_INSTRUCTIONS}\n\nTranscript:\n\"\"\"\n{transcript}\n\"\"\""

    with CLAUDE.messages.stream(
        model="claude-opus-4-7",
        max_tokens=8192,
        thinking={"type": "adaptive"},
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        msg = stream.get_final_message()

    text = "".join(b.text for b in msg.content if b.type == "text").strip()
    if text.startswith("```"):
        text = text.split("\n", 1)[1].rsplit("```", 1)[0].strip()
    rows = json.loads(text)
    return jsonify({"rows": rows})


@app.route("/api/docx", methods=["POST"])
@limiter.limit("30 per hour")
@limiter.limit("100 per day")
def docx():
    data = request.json
    rows = data["rows"]
    meta = data.get("meta", {})

    doc = Document()
    doc.add_heading(f"Bat Activity Log — {meta.get('site') or 'Site'}", level=1)

    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(f"{meta.get('date', '')}    ")
    p.add_run("Surveyor: ").bold = True
    p.add_run(f"{meta.get('surveyor', '')}    ")
    p.add_run("Survey type: ").bold = True
    p.add_run(f"{meta.get('survey_type', 'Dusk emergence')}")

    headers = ["Time", "Species", "Count", "Behaviour", "Location", "Direction", "Notes"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for r_idx, row in enumerate(rows, start=1):
        cells = [
            row.get("time", ""),
            row.get("species", ""),
            row.get("count", "1"),
            row.get("behaviour", ""),
            row.get("location", ""),
            row.get("direction", ""),
            row.get("notes", ""),
        ]
        for c_idx, val in enumerate(cells):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    widths_cm = [1.4, 3.2, 1.0, 2.2, 3.6, 1.2, 3.4]
    for col, w in zip(table.columns, widths_cm):
        for cell in col.cells:
            cell.width = Cm(w)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    site = meta.get("site") or "bat_survey"
    safe_site = "".join(c if c.isalnum() else "_" for c in site)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{safe_site}_activity_log.docx",
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)
